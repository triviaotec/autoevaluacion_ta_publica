# -*- coding: utf-8 -*-
"""
Autoevaluación Transparencia Activa · Versión cloud 2025
-----------------------------------------------------------------
• Optimizaciones para despliegue en Streamlit Cloud (caché, forms)
• Validaciones y reglas de negocio del script original restablecidas
• Mejoras de UX, reporte Word y detección dinámica de logo del script
  modificado conservadas

© 2025 – Diego González
"""
from __future__ import annotations
import base64, json
from collections import defaultdict, namedtuple
from datetime import datetime
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

# ──────────────────────────────── Paths ────────────────────────────────────
BASE   = Path(__file__).resolve().parent
P_ITEMS = BASE / "estructura_materias_items.json"
P_ESP   = BASE / "estructura_indicadores_especificos_ACTUALIZADO.json"
P_DOCX  = BASE / "plantilla_nueva.docx"
P_LOGO  = next((BASE / f"TRIVIA{ext}" for ext in (".png", ".jpeg", ".jpg") if (BASE / f"TRIVIA{ext}").exists()), None)

# ────────────────────────── Streamlit config  ──────────────────────────────
st.set_page_config(
    page_title="Autoevaluación Transparencia Activa",
    layout="wide",
    page_icon=str(P_LOGO) if P_LOGO else None,
)

def _logo64(p: Path | None) -> str:
    """Retorna el logo base‑64 o cadena vacía si no existe."""
    if not p or not p.exists():
        return ""
    return f"data:image/{p.suffix.lstrip('.')};base64," + base64.b64encode(p.read_bytes()).decode()

st.markdown(
    f"""
<style>
#MainMenu, header, footer{{visibility:hidden}}
#logo{{position:fixed;top:8px;right:18px;z-index:10}}
</style>
<div id="logo"><img src="{_logo64(P_LOGO)}" width="140"></div>
""",
    unsafe_allow_html=True,
)

_safe_rerun = lambda: (st.rerun if hasattr(st, "rerun") else st.experimental_rerun)()

# ──────────────────────── Carga estructura ────────────────────────────────
@st.cache_data(show_spinner=False)
def load_struct():
    df = pd.read_json(P_ITEMS)
    df["Materia"] = df["Materia"].replace(
        "Actos y resoluciones que tengas efectos sobre terceros",
        "Actos y resoluciones con efectos sobre terceros",
    )
    df = df.sort_values("ID")
    ind_esp = json.loads(P_ESP.read_text(encoding="utf-8"))
    return df, ind_esp

df, IND_ESP = load_struct()
ITEMS = df.to_dict("records")
TOTAL = len(ITEMS)

ITEM_TO_MAT  = {r["Ítem"]: r["Materia"].strip() for r in ITEMS}
MAT_TO_ITEMS = defaultdict(list)
for r in ITEMS:
    MAT_TO_ITEMS[r["Materia"].strip()].append(r["Ítem"])

ORDER_MAT: list[str] = []
_seen = set()
for r in ITEMS:                      # primera aparición → orden por ID
    m = r["Materia"].strip()
    if m not in _seen:
        ORDER_MAT.append(m); _seen.add(m)

def _num_first(s: pd.Series):
    s = pd.to_numeric(s, errors="coerce").dropna()
    return s.iloc[0] if not s.empty else pd.NA

MAT_PESO = df.groupby("Materia")["Peso Materia (%)"].apply(_num_first).to_dict()

# ───────────────────────────── US state ───────────────────────────────────
ItemR = namedtuple("ItemR", "scenario gen spec")  # gen=(disp, act, comp)
st.session_state.setdefault("idx", 0)
st.session_state.setdefault("answers", {})

# ──────────────────────────── Constantes ──────────────────────────────────
VAL_GEN  = {"Sí": 100, "No": 0, "No es posible determinarlo": 25}
GEN_DESC = {
    "disponibilidad": "Información no disponible",
    "actualización": "Información desactualizada",
    "completitud":   "Información incompleta",
}

# ----------------------------------------------------------------------------
# Utilidades
# ----------------------------------------------------------------------------

def _safe_idx(options: list[str], value, default: int = 0):
    """Índice seguro dentro de *options*."""
    try:
        return options.index(value) if value in options else default
    except (ValueError, TypeError):
        return default

# ----------------------------------------------------------------------------
# Cálculos de puntaje
# ----------------------------------------------------------------------------

def _puntaje_item(r: ItemR) -> float | None:
    """Retorna % de cumplimiento (0‑100) o None si se excluye."""
    if r.scenario in (2, 3):
        return None  # No aplica
    if r.scenario in (4, 5):
        return 0     # Infracción

    g_vals = [VAL_GEN[v] for v in r.gen if v is not None]
    if len(g_vals) < 3:
        return None  # Ítem incompleto
    gen_score = min(g_vals)

    esp_apl = [v for v in r.spec if v != "No aplica"]
    esp_score = 100 if not esp_apl else round(sum(v == "Sí" for v in esp_apl) / len(esp_apl) * 100)

    return round(gen_score * 0.75 + esp_score * 0.25, 1)


def _calcular() -> tuple[dict[str, float | None], dict[str, float | None], float]:
    """Calcula puntajes por ítem, materia y global."""
    item_sc = {it: _puntaje_item(r) for it, r in st.session_state.answers.items()}

    mat_sc: dict[str, float | None] = {}
    for m in ORDER_MAT:
        nums = [item_sc[i] for i in MAT_TO_ITEMS[m] if item_sc.get(i) is not None]
        mat_sc[m] = None if not nums else round(sum(nums) / len(nums), 1)

    pesos = {m: p for m, p in MAT_PESO.items() if pd.notna(p) and mat_sc.get(m) is not None}
    if pesos:
        glob = round(sum(mat_sc[m] * pesos[m] for m in pesos) / sum(pesos.values()), 1)
    else:
        vals = [v for v in mat_sc.values() if v is not None]
        glob = round(sum(vals) / len(vals), 1) if vals else 0
    return item_sc, mat_sc, glob

# ----------------------------------------------------------------------------
# Validación antes de guardar
# ----------------------------------------------------------------------------

def _valid_inputs(esc, disp, act, comp) -> bool:
    """Valida los campos obligatorios según reglas de negocio."""
    if esc != 1:
        return True  # Para escenarios 2‑5 no se requieren preguntas

    # Escenario 1
    if disp is None:
        return False
    if disp == "No":
        return True  # se cierra el ítem

    # Si hay información disponible…
    if act is None:
        return False
    if act == "No":
        return True  # se cierra aquí

    # Si además está actualizada…
    return comp is not None

# ----------------------------------------------------------------------------
# Tabla para Word (mejorada)
# ----------------------------------------------------------------------------

def _tabla(doc: Document, headers: list[str], rows: list[tuple[str, float | None]]):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.paragraphs[0].runs:
            r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)
        c._tc.get_or_add_tcPr().append(parse_xml(f"<w:shd {nsdecls('w')} w:fill='000000'/>") )

    # Filas de datos
    for r_i, (k, v) in enumerate(rows, 1):
        tbl.rows[r_i].cells[0].text = str(k)
        if v is None:
            # Distinción: no aplica vs no se evalúa
            txt = "No se evalúa"
            if isinstance(k, str) and k in st.session_state.answers:
                scen = st.session_state.answers[k].scenario
                if scen in (2, 3):
                    txt = "No aplica"
            tbl.rows[r_i].cells[1].text = txt
        else:
            tbl.rows[r_i].cells[1].text = f"{v:.1f}"
        tbl.rows[r_i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Bordes
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for side in ("top", "left", "bottom", "right"):
                if tcPr.find(qn(f"w:{side}")) is None:
                    tcPr.append(parse_xml(rf"<w:{side} w:val='single' w:sz='4' w:color='000000' {nsdecls('w')}/>"))

# ----------------------------------------------------------------------------
# Exportar Word
# ----------------------------------------------------------------------------

def _export(mat_sc, item_sc, glob, infr, org, evalua):
    doc = Document(str(P_DOCX))
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = s.right_margin = Cm(2)

    # Título
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REPORTE DE AUTOEVALUACIÓN DE TRANSPARENCIA ACTIVA"); run.bold = True; run.font.size = Pt(16)

    # Encabezado
    doc.add_paragraph(f"Organismo: {org}")
    doc.add_paragraph(f"Fecha: {datetime.now():%d-%m-%Y}")
    doc.add_paragraph(f"Evaluador(a): {evalua}")
    doc.add_paragraph(f"Cumplimiento TA global observado: {glob:.1f} %")

    # Puntajes
    doc.add_paragraph(); _tabla(doc, ["Materia", "%"], [(m, mat_sc[m]) for m in ORDER_MAT])
    doc.add_paragraph(); _tabla(doc, ["Ítem", "%"], [(it, item_sc.get(it)) for it in ITEM_TO_MAT])

    # Incumplimientos
    if infr:
        doc.add_paragraph()
        h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = h.add_run("Incumplimientos detectados"); rh.bold = True; rh.font.size = Pt(13)
        for m in ORDER_MAT:
            if m not in infr:
                continue
            pm = doc.add_paragraph(); rm = pm.add_run(m); rm.bold = True
            for it, lst in infr[m].items():
                pit = doc.add_paragraph(f"  {it}"); pit.runs[0].bold = True
                for i, txt in enumerate(lst, 1):
                    doc.add_paragraph(f"    {i}. {txt}")

    fname = f"Reporte_TA_{org}_{datetime.now():%Y%m%d}.docx"
    doc.save(BASE / fname)
    return fname

# ============================================================================
# INTERFAZ PRINCIPAL
# ============================================================================

cur = ITEMS[st.session_state.idx]
mat, it = cur["Materia"].strip(), cur["Ítem"]

st.title("Autoevaluación Transparencia Activa")
st.markdown(f"**Materia:** {mat}")
st.markdown(f"**Ítem:** {it}")

prev: ItemR | None = st.session_state.answers.get(it)
key = lambda s: f"{it}::{s}"

ESC_D = {
    1: "Organismo presenta sección con antecedentes",
    2: "Organismo indica no tener antecedentes / no aplica",
    3: "No hay sección pero no hay evidencia de infracción",
    4: "No hay sección y sí hay evidencia de información faltante",
    5: "Sección/vínculo existe pero no funciona / no muestra datos",
}
# ------------- FORM -------------
with st.form(key=f"form_{it}"):
    esc_idx = _safe_idx(list(ESC_D), prev.scenario if prev else None)
    esc = st.radio(
        "Escenario:",
        list(ESC_D),
        format_func=lambda v: f"Escenario {v}: {ESC_D[v]}",
        index=esc_idx,
        key=key("esc"),
    )

    disp = act = comp = None
    spec_vals: list[str] = []

    if esc == 1:
        # Indicadores generales
        disp = st.radio("1⃣ ¿Información disponible?", ["Sí", "No"],
                         index=_safe_idx(["Sí", "No"], prev.gen[0] if prev else None), key=key("disp"))

        if disp == "Sí":
            act = st.radio("2⃣ ¿Información actualizada?", ["Sí", "No"],
                            index=_safe_idx(["Sí", "No"], prev.gen[1] if prev else None), key=key("act"))
            if act == "Sí":
                opts_comp = ["Sí", "No", "No es posible determinarlo"]
                comp = st.radio("3⃣ ¿Información completa?", opts_comp,
                                index=_safe_idx(opts_comp, prev.gen[2] if (prev and len(prev.gen)>2) else None),
                                key=key("comp"))

        # Indicadores específicos
        if disp == act == "Sí" and comp is not None:
            lista = IND_ESP.get(f"{mat} || {it}", [])
            if lista:
                st.subheader("Indicadores específicos")
            for i, txt in enumerate(lista):
                radios = ["Sí", "No", "No aplica"]
                default = radios.index(prev.spec[i]) if (prev and i < len(prev.spec)) else 0
                spec_vals.append(st.radio(txt, radios, index=default, key=key(f"spec{i}")))

    submitted = st.form_submit_button("💾 Guardar ítem")

if submitted:
    if not _valid_inputs(esc, disp, act, comp):
        st.warning("Completa los indicadores antes de guardar.")
    else:
        st.session_state.answers[it] = ItemR(esc, (disp, act, comp), spec_vals)
        st.success("✔ Ítem guardado")

# ───────────────────────── Navegación ──────────────────────────────────────
col1, col2, col3 = st.columns(3)
with col1:
    if st.session_state.idx > 0 and st.button("⟵ Anterior"):
        st.session_state.idx -= 1; _safe_rerun()
with col2:
    st.markdown(f"**Ítem {st.session_state.idx + 1}/{TOTAL}**")
with col3:
    if st.session_state.idx < TOTAL - 1 and st.button("Siguiente ⟶"):
        st.session_state.idx += 1; _safe_rerun()

# ─────────────────────────── Sidebar ───────────────────────────────────────
st.sidebar.header("Resultados")
org_in  = st.sidebar.text_input("Nombre organismo")
eval_in = st.sidebar.text_input("Evaluador(a)")

if st.sidebar.button("Calcular resultados"):
    st.session_state.it_sc, st.session_state.mat_sc, st.session_state.glob_sc = _calcular()
    st.sidebar.metric("Cumplimiento global", f"{st.session_state.glob_sc:.1f} %")

st.sidebar.markdown("---")
if st.sidebar.button("Exportar Word") and org_in and eval_in:
    if "mat_sc" not in st.session_state:
        st.sidebar.warning("Calcula primero los resultados.")
    else:
        infr = defaultdict(lambda: defaultdict(list))
        for it, ans in st.session_state.answers.items():
            m = ITEM_TO_MAT[it]
            # Escenarios de infracción
            if ans.scenario in (4, 5):
                infr[m][it].append(ESC_D[ans.scenario])

            # Indicadores generales
            for idx, val in enumerate(ans.gen):
                if val in ("No", "No es posible determinarlo"):
                    infr[m][it].append(GEN_DESC[list(GEN_DESC)[idx]])

            # Indicadores específicos
            lista = IND_ESP.get(f"{m} || {it}", [])
            for idx, val in enumerate(ans.spec):
                if val == "No" and idx < len(lista):
                    infr[m][it].append(f"Indicador «{lista[idx]}» = No")

        fname = _export(
            st.session_state.mat_sc,
            st.session_state.it_sc,
            st.session_state.glob_sc,
            infr,
            org_in,
            eval_in,
        )
        with open(BASE / fname, "rb") as f:
            st.sidebar.download_button("📄 Descargar informe", f, file_name=fname)
        st.sidebar.success("Informe generado.")
